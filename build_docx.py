#!/usr/bin/env python3
"""Gera o Word do README com tudo em preto e tabelas bem formatadas."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

doc = Document()

# Estilos base — tudo preto
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
for h, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
    st = doc.styles[h]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = BLACK

TOKEN = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')

def add_rich(p, text):
    """Adiciona texto com **negrito**, `mono` e *itálico*, sempre em preto."""
    for seg in TOKEN.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith("`") and seg.endswith("`"):
            r = p.add_run(seg[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif seg.startswith("*") and seg.endswith("*"):
            r = p.add_run(seg[1:-1]); r.italic = True
        else:
            r = p.add_run(seg)
        r.font.color.rgb = BLACK

def para(text="", after=6, size=None):
    p = doc.add_paragraph()
    if text:
        add_rich(p, text)
    p.paragraph_format.space_after = Pt(after)
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    add_rich(p, text)
    p.paragraph_format.space_after = Pt(2)
    return p

def quote(text):
    p = doc.add_paragraph()
    add_rich(p, text)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.italic = True
    # barra à esquerda
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10"); left.set(qn("w:color"), "999999")
    pbdr.append(left); pPr.append(pbdr)
    return p

def code_block(lines):
    p = doc.add_paragraph()
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(10); r.font.color.rgb = BLACK
        if i < len(lines) - 1:
            r.add_break()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.1)
    return p

def set_cell(cell, text, bold=False, header=False):
    cell.text = ""
    par = cell.paragraphs[0]
    add_rich(par, text)
    for r in par.runs:
        r.bold = bold
        r.font.size = Pt(9.5)
        r.font.color.rgb = BLACK
    if header:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "D9E2F3")
        tcPr.append(shd)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, header=True)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            set_cell(cells[i], c)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ===== Conteúdo =====
h = doc.add_heading("Lab05 — Desenho e Preparação do Experimento", level=1)
doc.add_heading("GraphQL vs REST — Um experimento controlado", level=2)
para("**Aluno:** Pedro Lucas Sousa e Silva")
quote("Este documento descreve o desenho (Sprint 1) e a preparação do experimento, já "
      "alinhado com o que foi efetivamente utilizado na execução. A execução, a análise "
      "estatística e a discussão dos resultados estão em RELATORIO_FINAL.md.")

doc.add_heading("1. Desenho do Experimento", level=2)

doc.add_heading("A. Hipóteses", level=3)
para("**Hipótese nula (H0):** Não existe diferença significativa no tempo de resposta nem no tamanho das respostas entre consultas GraphQL e REST.")
para("**Hipótese alternativa (H1):** Consultas GraphQL apresentam tempo de resposta menor e tamanho das respostas menor quando comparadas às consultas REST.")
quote("Observação: por H1 ser unicaudal (\"menor\"), os testes estatísticos foram aplicados na forma one-sided (t pareado e Wilcoxon unicaudal).")
quote("Resultado (resumo): a hipótese se confirmou para o tamanho (GraphQL ~90% menor) e foi refutada para o tempo (GraphQL ficou mais lento que REST). Detalhes no relatório final.")

doc.add_heading("B. Variáveis Dependentes", level=3)
para("São as métricas medidas:")
bullet("Tempo de resposta (ms)")
bullet("Tamanho da resposta (bytes)")

doc.add_heading("C. Variáveis Independentes", level=3)
para("Fatores manipulados no experimento:")
bullet("Tipo de API: REST vs GraphQL")
bullet("Tipo de consulta: simples, filtrada, por ID e composta")

doc.add_heading("D. Tratamentos", level=3)
bullet("**T1 — REST:** consultas executadas usando endpoints REST tradicionais.")
bullet("**T2 — GraphQL:** as mesmas consultas executadas usando sintaxe GraphQL equivalente.")

doc.add_heading("E. Objetos Experimentais", level=3)
para("Foi utilizada a **API pública do Rick and Morty**, que oferece as duas abordagens sobre os mesmos dados e **não exige token/autenticação**:")
bullet("REST: `https://rickandmortyapi.com/api`")
bullet("GraphQL: `https://rickandmortyapi.com/graphql`")
para("Sobre ela são feitas 4 consultas equivalentes:")
bullet("consulta simples — lista de personagens;")
bullet("consulta filtrada — personagens com `name=rick`;")
bullet("consulta por ID — um personagem específico;")
bullet("consulta composta — personagem + episódios relacionados.")

doc.add_heading("F. Tipo de Projeto Experimental", level=3)
para("Projeto controlado com **medidas repetidas (within-subject / pareado)**, pois:")
bullet("os mesmos objetos são testados em REST e em GraphQL;")
bullet("o mesmo ambiente é usado para evitar ruído externo;")
bullet("a comparação é direta entre tratamentos (pareada por consulta e repetição).")

doc.add_heading("G. Quantidade de Medições", level=3)
para("Para cada consulta:")
bullet("30 medições em REST;")
bullet("30 medições em GraphQL (+ 1 repetição de *warm-up* descartada por consulta).")
para("**Total:** 4 consultas × 30 repetições × 2 tratamentos = **240 medições** (120 pares).")

doc.add_heading("H. Ameaças à Validade", level=3)
para("**Validade interna**", after=2)
bullet("Oscilações de rede podem afetar tempos → mesmo computador e rede; repetições e uso da mediana.")
bullet("**Cache de CDN (observado):** as respostas REST (`GET`) são servidas do cache de borda do CDN (Cloudflare) e voltam muito rápidas, enquanto o GraphQL (`POST`) não é cacheável e sempre atinge o servidor de origem. Isso favorece o REST no tempo. Mitigação parcial: cabeçalho `Cache-Control: no-cache` e warm-up; o efeito é registrado como limitação e discutido no relatório.")
para("**Validade externa**", after=2)
bullet("Uso de apenas um domínio (Rick and Morty) → os resultados podem não generalizar para outras APIs. Limitação assumida.")
para("**Validade de construção**", after=2)
bullet("REST e GraphQL solicitam o mesmo conjunto de campos úteis. O REST, porém, retorna campos extras (*over-fetching*), o que faz parte do fenômeno estudado e explica boa parte da diferença de tamanho.")
bullet("Tempo medido como latência fim-a-fim da requisição bem-sucedida (as esperas de retentativa por rate limit são excluídas da medição).")
para("**Validade estatística / de conclusão**", after=2)
bullet("30 repetições por consulta mitigam ruído; além do p-valor, reporta-se o tamanho de efeito (Cohen d e r). Como a normalidade não é assumida, aplicam-se dois testes (t pareado e Wilcoxon).")

doc.add_heading("2. Preparação do Experimento", level=2)

doc.add_heading("2.1 Ambiente", level=3)
bullet("Computador pessoal (macOS — MacBook Air).")
bullet("**Python 3.14** como linguagem principal.")
bullet("Bibliotecas: `requests` (chamadas REST e GraphQL, com sessão keep-alive e retentativas com backoff); `pandas` e `numpy` (análise); `csv` e `time.perf_counter` (padrão).")
bullet("**Testes estatísticos:** t de Student pareado e Wilcoxon **implementados em Python puro** (sem `scipy`).")
bullet("**Visualização:** gráficos HTML interativos com **Chart.js**.")
quote("Decisão de ferramenta: optou-se por Python + requests por integração direta com pandas/numpy na análise, mantendo todo o pipeline (coleta → análise → gráficos) em uma só linguagem.")

doc.add_heading("2.2 Scripts", level=3)
table(["Arquivo", "Função"],
      [["experimento.py", "Coleta real na API do Rick and Morty; mede tempo e tamanho; salva resultados.csv."],
       ["gerar_dados.py", "Gera um resultados.csv simulado (fallback para testar a pipeline sem coletar)."],
       ["analise.py", "Estatística pareada (t + Wilcoxon), tamanho de efeito; gera analise_resumo.json."],
       ["gerar_graficos.py", "Gera os 2 gráficos HTML em graficos/ (tempo e tamanho)."]],
      [1.7, 4.8])
para("Ordem de execução:")
code_block([
    "pip install -r requirements.txt",
    "python3 experimento.py        # coleta os dados reais",
    "python3 analise.py            # roda a estatística",
    "python3 gerar_graficos.py     # gera os gráficos",
])

doc.add_heading("2.3 Consultas", level=3)
para("As 4 consultas, equivalentes em REST e GraphQL:")
table(["#", "Consulta", "REST", "GraphQL"],
      [["1", "Simples", "GET /character", "characters { results { … } }"],
       ["2", "Filtrada", "GET /character/?name=rick", "characters(filter:{name:\"rick\"})"],
       ["3", "Por ID", "GET /character/1", "character(id:1)"],
       ["4", "Composta", "GET /character/1 + GET /episode/1,2,… (lote)", "character(id:1){ … episode{…} }"]],
      [0.4, 1.0, 2.7, 2.4])

doc.add_heading("2.4 Estrutura do CSV", level=3)
para("Cada linha registrada segue o formato:")
table(["metodo", "consulta", "repeticao", "tempo_ms", "tamanho_bytes"],
      [["REST", "simples", "1", "56.2", "19496"],
       ["GraphQL", "simples", "1", "321.2", "3421"]],
      [1.3, 1.3, 1.35, 1.25, 1.3])
para("Isso permite as análises diretas para RQ1 (tempo) e RQ2 (tamanho).")

# Imagens (gráficos)
for img in ["image-1.png", "image.png"]:
    try:
        doc.add_picture(img, width=Inches(6.2))
    except Exception as e:
        para(f"[imagem {img} não encontrada: {e}]")

doc.save("Lab05_Desenho_Preparacao.docx")
print("DOCX_OK")
